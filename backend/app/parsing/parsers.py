import csv
from datetime import date, datetime, time
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from .models import ParsedDocument, ParsedDocumentSection


TEXT_EXTENSIONS = {".md", ".txt"}
CSV_EXTENSIONS = {".csv"}
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx"}


class UnsupportedDocumentTypeError(ValueError):
    pass


class EmptyParsedDocumentError(ValueError):
    pass


class UnreadableDocumentError(ValueError):
    pass


def parse_file(file_path: Path, file_extension: str) -> ParsedDocument:
    normalized_extension = file_extension.lower()
    if normalized_extension in TEXT_EXTENSIONS:
        return parse_text_file(file_path)
    if normalized_extension in CSV_EXTENSIONS:
        return parse_csv_file(file_path)
    if normalized_extension in PDF_EXTENSIONS:
        return parse_pdf_file(file_path)
    if normalized_extension in DOCX_EXTENSIONS:
        return parse_docx_file(file_path)
    if normalized_extension in XLSX_EXTENSIONS:
        return parse_xlsx_file(file_path)

    raise UnsupportedDocumentTypeError(f"Parsing is not supported for {file_extension} files yet")


def parse_text_file(file_path: Path) -> ParsedDocument:
    text = file_path.read_text(encoding="utf-8-sig")
    normalized_text = normalize_text(text)

    return ParsedDocument(
        sections=[
            ParsedDocumentSection(
                kind="text",
                label="Document text",
                text=normalized_text,
                metadata={"parser": "text"},
            )
        ]
    )


def parse_csv_file(file_path: Path) -> ParsedDocument:
    with file_path.open("r", encoding="utf-8-sig", newline="") as csv_file:
        rows = list(csv.reader(csv_file))

    text = format_csv_rows(rows)
    row_count = len(rows)
    column_count = max((len(row) for row in rows), default=0)

    return ParsedDocument(
        sections=[
            ParsedDocumentSection(
                kind="table",
                label="CSV table",
                text=text,
                metadata={
                    "parser": "csv",
                    "row_count": row_count,
                    "column_count": column_count,
                },
            )
        ]
    )


def format_csv_rows(rows: list[list[str]]) -> str:
    return "\n".join(
        f"Row {row_index}: {format_csv_row(row)}"
        for row_index, row in enumerate(rows, start=1)
    ).strip()


def format_csv_row(row: list[str]) -> str:
    return " | ".join(normalize_text(cell) for cell in row)


def parse_pdf_file(file_path: Path) -> ParsedDocument:
    try:
        reader = PdfReader(file_path)
        page_count = len(reader.pages)
    except Exception as error:
        raise UnreadableDocumentError("PDF could not be read") from error

    sections: list[ParsedDocumentSection] = []

    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = normalize_text(page.extract_text() or "")
        except Exception as error:
            raise UnreadableDocumentError(f"Text could not be extracted from page {page_index}") from error

        if not text:
            continue

        sections.append(
            ParsedDocumentSection(
                kind="page",
                label=f"Page {page_index}",
                text=text,
                metadata={
                    "parser": "pypdf",
                    "page_number": page_index,
                    "page_count": page_count,
                },
            )
        )

    if not sections:
        raise EmptyParsedDocumentError("No extractable text found in PDF")

    return ParsedDocument(sections=sections)


def parse_docx_file(file_path: Path) -> ParsedDocument:
    try:
        document = DocxDocument(file_path)
    except Exception as error:
        raise UnreadableDocumentError("DOCX file could not be read") from error

    paragraphs = [
        normalize_text(paragraph.text)
        for paragraph in document.paragraphs
        if normalize_text(paragraph.text)
    ]
    table_rows = extract_docx_table_rows(document)
    text_parts = paragraphs + table_rows
    text = "\n\n".join(text_parts)

    if not text:
        raise EmptyParsedDocumentError("No extractable text found in DOCX file")

    return ParsedDocument(
        sections=[
            ParsedDocumentSection(
                kind="text",
                label="Word document",
                text=text,
                metadata={
                    "parser": "python-docx",
                    "paragraph_count": len(paragraphs),
                    "table_count": len(document.tables),
                    "table_row_count": len(table_rows),
                },
            )
        ]
    )


def extract_docx_table_rows(document) -> list[str]:
    rows: list[str] = []

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [normalize_text(cell.text) for cell in row.cells]
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                rows.append(f"Table {table_index}, row {row_index}: {row_text}")

    return rows


def parse_xlsx_file(file_path: Path) -> ParsedDocument:
    try:
        workbook = load_workbook(file_path, read_only=True, data_only=True)
    except Exception as error:
        raise UnreadableDocumentError("XLSX file could not be read") from error

    sections: list[ParsedDocumentSection] = []
    for sheet in workbook.worksheets:
        rows = extract_xlsx_rows(sheet)
        if not rows:
            continue

        row_count = len(rows)
        column_count = max((len(row) for row in rows), default=0)
        sections.append(
            ParsedDocumentSection(
                kind="sheet",
                label=f"Sheet: {sheet.title}",
                text=format_xlsx_rows(rows),
                metadata={
                    "parser": "openpyxl",
                    "sheet_name": sheet.title,
                    "row_count": row_count,
                    "column_count": column_count,
                },
            )
        )

    workbook.close()

    if not sections:
        raise EmptyParsedDocumentError("No extractable text found in XLSX file")

    return ParsedDocument(sections=sections)


def extract_xlsx_rows(sheet) -> list[list[str]]:
    rows: list[list[str]] = []

    for row in sheet.iter_rows(values_only=True):
        formatted_cells = [format_cell_value(value) for value in row]
        while formatted_cells and not formatted_cells[-1]:
            formatted_cells.pop()

        if any(formatted_cells):
            rows.append(formatted_cells)

    return rows


def format_xlsx_rows(rows: list[list[str]]) -> str:
    return "\n".join(
        f"Row {row_index}: {format_xlsx_row(row)}"
        for row_index, row in enumerate(rows, start=1)
    ).strip()


def format_xlsx_row(row: list[str]) -> str:
    return " | ".join(row)


def format_cell_value(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.isoformat(sep=" ")
    if isinstance(value, (date, time)):
        return value.isoformat()

    return normalize_text(str(value))


def normalize_text(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n").strip()
